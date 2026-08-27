import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os

def set_cell_background(cell, fill_color):
    """Sets background color of a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tc_pr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Sets cell padding."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tc_pr.append(tc_mar)

def set_rtl(paragraph_or_run):
    """Marks paragraph or run as RTL."""
    p_pr = paragraph_or_run._element.get_or_add_pPr() if hasattr(paragraph_or_run, '_element') and paragraph_or_run._element.tag.endswith('p') else None
    if p_pr is not None:
        bidi = parse_xml(f'<w:bidi {nsdecls("w")}/>')
        p_pr.append(bidi)

def create_document():
    doc = Document()
    
    # Page setup - A4 with 1 inch margins
    for section in doc.sections:
        # Set margins
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        
        # Set RTL for section if possible
        try:
            sectPr = section._sectPr
            bidi = parse_xml(f'<w:bidi {nsdecls("w")}/>')
            sectPr.append(bidi)
        except Exception:
            pass

    # Styles & Colors
    PRIMARY = RGBColor(13, 148, 136)     # #0d9488 Teal
    SECONDARY = RGBColor(15, 118, 110)   # #0f766e Dark Teal
    DARK = RGBColor(30, 41, 59)          # #1e293b Slate Dark
    GRAY = RGBColor(100, 116, 139)       # #64748b Gray
    MUTED = RGBColor(71, 85, 105)

    def add_title_page():
        p_univ = doc.add_paragraph()
        p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(p_univ)
        r = p_univ.add_run("الجمهورية العربية السورية\nوزارة التعليم العالي والبحث العلمي\nكلية الهندسة المعلوماتية / علوم الحاسوب\nقسم هندسة البرمجيات ونظم المعلومات\n")
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = DARK

        # Spacing
        for _ in range(2):
            doc.add_paragraph()

        # Project Title Box
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(p_title)
        r_sub = p_title.add_run("مشروع تخرج لنيل درجة الإجازة في الهندسة المعلوماتية\n\n")
        r_sub.font.name = 'Arial'
        r_sub.font.size = Pt(13)
        r_sub.font.color.rgb = GRAY

        r_title = p_title.add_run("نظام إدارة النقاط والمراكز الطبية المؤتمت\n(Medical Point Management System)\n")
        r_title.font.name = 'Arial'
        r_title.font.size = Pt(22)
        r_title.font.bold = True
        r_title.font.color.rgb = PRIMARY

        r_desc = p_title.add_run("\nنظام ويب إكلينيكي متكامل لرقمنة رحلة المريض، الاستقبال، الفحوصات المخبرية، والصيدلية\n")
        r_desc.font.name = 'Arial'
        r_desc.font.size = Pt(13)
        r_desc.font.italic = True
        r_desc.font.color.rgb = SECONDARY

        for _ in range(3):
            doc.add_paragraph()

        # Info Table
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        cell_r0_c0 = table.cell(0, 0)
        cell_r0_c1 = table.cell(0, 1)
        cell_r1_c0 = table.cell(1, 0)
        cell_r1_c1 = table.cell(1, 1)

        def style_info_cell(cell, label, val):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            r1 = p.add_run(f"{label}\n")
            r1.font.name = 'Arial'
            r1.font.size = Pt(12)
            r1.font.bold = True
            r1.font.color.rgb = PRIMARY
            r2 = p.add_run(val)
            r2.font.name = 'Arial'
            r2.font.size = Pt(12)
            r2.font.color.rgb = DARK

        style_info_cell(cell_r0_c1, "إعداد الطالب / فريق العمل:", "• مهند الغزالي\n• فريق العمل المطور")
        style_info_cell(cell_r0_c0, "إشراف الدكتور / الأستاذ:", "• أستاذ المادة والمشرف الأكاديمي")
        style_info_cell(cell_r1_c1, "التخصص الأكاديمي:", "هندسة البرمجيات ونظم المعلومات الصحية")
        style_info_cell(cell_r1_c0, "العام الجامعي:", "2025 - 2026 م")

        for row in table.rows:
            for cell in row.cells:
                set_cell_background(cell, "F8FAFC")
                set_cell_margins(cell, 150, 150, 200, 200)

        doc.add_page_break()

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = SECONDARY
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = DARK
        return p

    def add_body_p(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_rtl(p)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(6)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.name = 'Arial'
            r_pre.font.size = Pt(11)
            r_pre.font.bold = True
            r_pre.font.color.rgb = DARK
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.font.italic = italic
        r.font.color.rgb = DARK
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.name = 'Arial'
            r_pre.font.size = Pt(11)
            r_pre.font.bold = True
            r_pre.font.color.rgb = SECONDARY
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        return p

    def add_callout(title, content):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "F0FDFA") # Light teal
        set_cell_margins(cell, 150, 150, 200, 200)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        r_t = p.add_run(f"💡 {title}\n")
        r_t.font.name = 'Arial'
        r_t.font.size = Pt(11.5)
        r_t.font.bold = True
        r_t.font.color.rgb = SECONDARY
        r_c = p.add_run(content)
        r_c.font.name = 'Arial'
        r_c.font.size = Pt(10.5)
        r_c.font.color.rgb = DARK
        doc.add_paragraph() # space

    def add_styled_table(headers, rows_data):
        table = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Style Header
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            set_cell_background(cell, "0D9488")
            set_cell_margins(cell, 140, 140, 150, 150)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(p)
            r = p.add_run(h)
            r.font.name = 'Arial'
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

        # Style Data Rows
        for r_idx, row in enumerate(rows_data):
            bg = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
            for c_idx, val in enumerate(row):
                cell = table.cell(r_idx+1, c_idx)
                set_cell_background(cell, bg)
                set_cell_margins(cell, 100, 100, 120, 120)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_rtl(p)
                r = p.add_run(str(val))
                r.font.name = 'Arial'
                r.font.size = Pt(10)
                r.font.color.rgb = DARK
        doc.add_paragraph()

    # --- Build Document Content ---
    add_title_page()

    # 1. الإهداء والملخص
    add_heading_1("الإهداء والشكر والتقدير")
    add_body_p("نهدي هذا الجهد المتواضع إلى أسرنا الكريمة التي كانت وما زالت السند الحقيقي وراء كل خطوة نجاح، وإلى أساتذتنا الأفاضل في كلية الهندسة المعلوماتية الذين غرسوا فينا حب المعرفة وقدموا لنا كل العون العلمي والتوجيه السديد.")
    add_body_p("كما نتوجه بجزيل الشكر وعظيم الامتنان إلى مشرف المشروع على توجيهاته القيمة ومتابعته المستمرة التي أثرت هذا العمل وساهمت في خروجه بأفضل صورة برمجية وأكاديمية.")

    add_heading_1("الملخص التنفيذي (Abstract)")
    add_body_p("يهدف مشروع «نظام إدارة النقاط والمراكز الطبية (Medical Point Management System)» إلى رقمنة وأتمتة كافة العمليات الإدارية والسريرية داخل مراكز الرعاية الصحية الأولية والنقاط الطبية المستقلة. يعالج النظام مشكلة البيروقراطية والسجلات الورقية التقليدية وما ينجم عنها من هدر للوقت، ضياع للملفات، وصعوبة في متابعة مخزون الأدوية وتاريخ المرضى.")
    add_body_p("يقدم النظام حلاً برمجياً متكاملاً مبنياً وفق معمارية الويب الحديثة (Web-Based Application) يربط أربعة أدوار رئيسية بصلاحيات صارمة (Role-Based Access Control): مدير النظام، الطبيب المعالج، فني المختبر، والصيدلي. يشمل النظام دورة علاجية مؤتمتة تبدأ من تسجيل المريض في الاستقبال، مروراً بالكشف السريري وتدوين التشخيص وطلب التحاليل المخبرية إلكترونياً، ثم إدخال النتائج ومقارنتها بالمدى المرجعي، وصولاً إلى تحويل الوصفة الطبية لصيدلية المركز وصرف الدواء مع خصم فوري وتلقائي من المخزون وتنبيهات الرصيد.")

    # 2. الفصل الأول
    add_heading_1("الفصل الأول: الإطار العام للمشروع (Introduction & Context)")
    add_heading_2("1.1 مقدمة عامة وخلفية المشروع")
    add_body_p("شهدت الأنظمة الصحية خلال السنوات الأخيرة تحولاً جذرياً نحو الرقمنة الشاملة (Digital Healthcare Transformation)، نظراً لما توفره من دقة متناهية في حفظ البيانات وسرعة فائقة في اتخاذ القرارات العلاجية. وتعتبر مراكز الرعاية الصحية الأولية والنقاط الطبية خط الدفاع الأول في المنظومة الصحية، إلا أنها في كثير من الأحيان تعاني من الاعتماد على الدفاتر الورقية أو برمجيات منفصلة غير مترابطة، مما يسبب انقطاعاً في تدفق البيانات بين الطبيب والمختبر والصيدلية.")
    
    add_heading_2("1.2 مشكلة المشروع (Problem Statement)")
    add_body_p("تتلخص المشكلات الأساسية التي يعالجها المشروع في النقاط التالية:")
    add_bullet(" اعتماد السجلات والملفات الورقية المعرضة للتلف والضياع وصعوبة البحث في التاريخ المرضي السابق للمريض.", "1. هدر الوقت وتلف السجلات:")
    add_bullet(" نقل المريض للتحاليل المخبرية والوصفات الطبية ورقياً، مما يؤدي لأخطاء قراءة الخط الطبي أو فقدان النتائج.", "2. انقطاع التواصل الإكلينيكي:")
    add_bullet(" عدم وجود نظام صرف مباشر يخصم الكميات من مستودع الصيدلية، مما يسبب نفاذ الأدوية دون سابق إنذار.", "3. سوء إدارة المخزون الدوائي:")
    add_bullet(" عدم وجود واجهات إحصائية مركزية تعطي إدارة المركز تقارير لحظية عن عدد الزيارات، التحاليل المعلقة، والضغط اليومي.", "4. غياب الرؤية الإحصائية:")

    add_heading_2("1.3 أهداف المشروع (Project Objectives)")
    add_bullet("بناء نظام ويب إكلينيكي مركزي مؤتمت بالكامل يربط جميع محطات المركز الطبي.")
    add_bullet("تطبيق معايير السجل الطبي الإلكتروني الموحد (Unified Electronic Health Record) لكل مريض.")
    add_bullet("تأمين التواصل اللحظي الفوري بين عيادة الطبيب ومختبر التحاليل وصيدلية المركز.")
    add_bullet("أتمتة إدارة المستودع الدوائي مع الخصم الآلي عند الصرف والتنبيه عند هبوط الرصيد.")
    add_bullet("تطبيق أعلى معايير أمن المعلومات والتحكم في الوصول المبني على الأدوار (RBAC).")

    # 3. الفصل الثاني
    add_heading_1("الفصل الثاني: تحليل المتطلبات ونمذجة النظام (Requirements Analysis)")
    add_heading_2("2.1 المتطلبات الوظيفية (Functional Requirements)")
    add_body_p("تم تقسيم المتطلبات الوظيفية بدقة حسب الصلاحيات والأدوار الأربعة للنظام:")

    headers_rbac = ["الدور (Role)", "المسؤوليات والصلاحيات الأساسية", "الصفحة الافتراضية"]
    data_rbac = [
        ["مدير النظام (Admin)", "إدارة الطاقم الطبي والمستخدمين، إضافة وتعديل المحافظات والمدن والأحياء، إدارة أنواع التحاليل ودليل الأدوية، والاطلاع على لوحة الإحصائيات المركزية والتقارير الشاملة.", "/dashboard"],
        ["الطبيب (Doctor)", "معاينة قائمة الزيارات، فتح محطة الكشف الطبي، تدوين العلامات الحيوية والتشخيص، طلب فحوصات مخبرية بضغطة زر، الاطلاع على النتائج الحية ومقارنتها بالمدى الطبيعي، كتابة الوصفات الطبية، وطباعة الروشتة والسجل الطبي للمريض.", "/dashboard/visits"],
        ["فني المختبر (Lab Tech)", "استقبال طلبات الفحوصات المحولة من الأطباء لحظياً، إدخال قيم النتائج الرقمية والوصفية، التحقق الذكي من المدى المرجعي (طبيعي / شاذ)، واعتماد النتيجة وطباعة تقرير التحليل المخبري.", "/dashboard/lab-results"],
        ["الصيدلي (Pharmacist)", "استعراض الوصفات الطبية المعلقة المحولة من عيادات الأطباء، مراجعة الجرعات والتعليمات، صرف الدواء بنقرة واحدة مع خصم الكمية تلقائياً من المخزون، وإدارة مستويات الأدوية.", "/dashboard/pharmacy"]
    ]
    add_styled_table(headers_rbac, data_rbac)

    add_heading_2("2.2 المتطلبات غير الوظيفية (Non-Functional Requirements)")
    add_bullet("حماية كاملة لكافة واجهات الـ REST APIs باستخدام Sanctum Bearer Tokens، وحظر أي استدعاء غير مصرح به بكود 403 Forbidden.", "• الأمان والسرية (Security):")
    add_bullet("استجابة فورية للطلبات بزمن معالجة يقل عن 200ms، مع دعم التحديث الدوري الذكي (Auto-Polling) لتحديث البيانات بين الأقسام دون ريفريش يدوي.", "• الأداء والسرعة (Performance):")
    add_bullet("تصميم واجهات عربية تفاعلية حديثة ومتجاوبة (Responsive RTL UI) تعتمد أفضل معايير تجربة المستخدم (UX/UI).", "• سهولة الاستخدام (Usability):")
    add_bullet("بناء معمارية برمجية قابلة للتوسع واستيعاب آلاف السجلات الطبية والزيارات اليومية بسهولة.", "• القابلية للتوسع (Scalability):")

    add_heading_2("2.3 المخطط العام لحالات الاستخدام (Use Case Overview)")
    add_body_p("يوضح النظام مجموعة من سيناريوهات الاستخدام المترابطة:")
    add_bullet("تسجيل الدخول -> التحقق من الصلاحية -> التوجيه للمحطة المخصصة.")
    add_bullet("استقبال مريض -> تسجيل البيانات الديموغرافية والجغرافية -> فتح زيارة طبية.")
    add_bullet("فحص الطبيب -> تسجيل التشخيص -> طلب تحليل -> إدخال المخبر للنتيجة -> مراجعة الطبيب -> وصف دواء -> صرف الصيدلي.")

    # 4. الفصل الثالث
    add_heading_1("الفصل الثالث: التصميم المعماري وهندسة البيانات (Architecture & Database)")
    add_heading_2("3.1 المعمارية العامة للنظام (System Architecture)")
    add_body_p("يعتمد النظام على المعمارية الطبقية المترابطة (Layered MVC Architecture) مقترنة بنموذج خادم-عميل عبر واجهات البرمجة التطبيقية (RESTful API Engine):")
    add_bullet("طبقة الواجهات (Presentation Layer): صفحات Blade تفاعلية مدعمة بـ JavaScript ES6 و CSS مخصص بالكامل يدعم الاتجاه RTL والتجاوب مع كافة الشاشات.")
    add_bullet("طبقة التحكم والمنطق (Application & Business Logic Layer): مبنية بإطار العمل Laravel 10 وتتضمن الـ Controllers والـ Middlewares مثل CheckRole و Authenticate.")
    add_bullet("طبقة البيانات والوصول (Data Access & Persistence Layer): محرك MySQL مع الـ Eloquent ORM لتأمين العلاقات والترابط السلس بين الجداول والمعاملات الآمنة (DB Transactions).")

    add_heading_2("3.2 معجم وقاموس قاعدة البيانات (Data Dictionary)")
    add_body_p("يوضح الجدول التالي البنية الهيكلية الشاملة لجداول قاعدة البيانات والعلاقات الرابطة بينها:")

    headers_db = ["اسم الجدول", "الوصف والهدف", "أهم الحقول والمفاتيح", "العلاقات والارتباطات"]
    data_db = [
        ["users", "حسابات مستخدمي النظام وموظفي المركز", "id, email, password, role, person_id", "belongsTo(Person)"],
        ["persons", "السجل الديموغرافي للمرضى والموظفين", "id, full_name, national_id, phone, birth_date, gender, neighborhood_id", "hasMany(Visits), belongsTo(Neighborhood)"],
        ["visits", "جلسات وزيارات الكشف الطبي للمرضى", "id, person_id, doctor_id, appointment_date, status, diagnosis, doctor_notes, blood_pressure, weight, temperature", "belongsTo(Person), belongsTo(Doctor), hasMany(LabResults), hasMany(PrescriptionItems)"],
        ["test_types", "دليل وأنواع الفحوصات المخبرية والمدى المرجعي", "id, name, code, unit, min_range, max_range, price, status", "hasMany(LabResults)"],
        ["lab_results", "طلبات ونتائج الفحوصات المخبرية للمرضى", "id, visit_id, test_type_id, result_value, lab_notes, status", "belongsTo(Visit), belongsTo(TestType)"],
        ["medicines", "دليل ومستودع الأدوية والمخزون", "id, name, strength, stock_quantity, is_available", "hasMany(PrescriptionItems)"],
        ["prescription_items", "بنود الوصفات الطبية المصروفة والمعلقة", "id, visit_id, medicine_id, prescribed_quantity, instructions, is_dispensed, dispensed_at", "belongsTo(Visit), belongsTo(Medicine)"],
        ["governorates", "المحافظات الجغرافية", "id, name, status", "hasMany(Cities)"],
        ["cities", "المدن التابعة للمحافظات", "id, name, governorate_id, status", "belongsTo(Governorate), hasMany(Neighborhoods)"],
        ["neighborhoods", "الأحياء السكنية التابعة للمدن", "id, name, city_id, status", "belongsTo(City), hasMany(Persons)"]
    ]
    add_styled_table(headers_db, data_db)

    add_callout("نمط المعاملات البنكية في الصرف (Atomic Dispense Transaction)", 
                "تعتمد دالة الصرف PrescriptionController::dispense على تقنية DB::transaction البرمجية، حيث يتم التحقق من توفر الكمية بالمستودع أولاً، ثم خصم الرصيد وتحديث حالة البند إلى (تم الصرف) في خطوة ذرية غير قابلة للتجزئة تضمن سلامة واتساق البيانات المخزنية بنسبة 100%.")

    # 5. الفصل الرابع
    add_heading_1("الفصل الرابع: التنفيذ والواجهات البرمجية (Implementation & UI)")
    add_heading_2("4.1 التقنيات وأدوات التطوير المستخدمة (Tech Stack)")
    add_bullet("Laravel 10.x كإطار عمل متين وموثوق يوفر بنية أمنية قوية ومحرك Eloquent ORM المتطور.", "• بيئة الخادم (Backend):")
    add_bullet("قاعدة بيانات علائقية سريعة ومثالية لإدارة السجلات الطبية الضخمة مع دعم الترميز العربي utf8mb4.", "• قاعدة البيانات (Database):")
    add_bullet("نظام المصادقة وحماية الـ APIs باستخدام Sanctum Token Authentication المشفرة.", "• الأمان والتوثيق (Auth):")
    add_bullet("تصميم واجهات احترافية مخصصة بدون أطر ثقيلة لضمان أعلى سرعة تحميل مع تجاوب كامل للأجهزة اللوحية والمكتبية.", "• الواجهات الأمامية (Frontend):")

    add_heading_2("4.2 استعراض الشاشات والوحدات الرئيسية")
    add_bullet("تتضمن مؤشرات إحصائية لحظية لعدد المرضى، زيارات اليوم، الحالات في الانتظار، التحاليل والوصفات المعلقة، والأدوية منخفضة المخزون مع جداول وصول سريعة.", "1. لوحة التحكم والإحصائيات (Dashboard):")
    add_bullet("شاشة مدمجة تتيح للطبيب تدوين التشخيص والعلامات الحيوية، طلب تحاليل بضغطة زر، متابعة نتائج المختبر الحية مع تلوين النتيجة (أخضر=طبيعي، أحمر=شاذ)، وكتابة وطباعة الوصفة الطبية.", "2. محطة الكشف الطبي (Consultation Station):")
    add_bullet("شاشة مخصصة لفني المختبر تستعرض الطلبات الواردة فوراً، ونافذة إدخال ذكية تفحص القيمة وتقارنها بالمدى المرجعي، مع طباعة تقرير الفحص المعتمد.", "3. وحدة المختبر والفحوصات (Laboratory Workspace):")
    add_bullet("شاشة تفاعلية تستعرض الوصفات المحولة من الأطباء، تتيح للصيدلي التحقق من تعليمات الطبيب وصرف الدواء بنقرة واحدة مع الخصم التلقائي من المخزون.", "4. وحدة الصيدلية وصرف الأدوية (Pharmacy Console):")
    add_bullet("سجل كامل لبيانات المرضى يتيح استعراض التاريخ المرضي والملف الطبي التراكمي (Timeline) بضغطة زر واحدة.", "5. سجل المرضى والتاريخ الطبي (Patient Registry & History):")

    # 6. الفصل الخامس
    add_heading_1("الفصل الخامس: الاختبارات وضمان الجودة (Testing & Verification)")
    add_heading_2("5.1 منهجية الاختبار (Testing Methodology)")
    add_body_p("تم إخضاع النظام لاختبارات صارمة شملت اختبارات التكامل (Integration Tests) واختبارات الأمان واختبار دورة العمل الإكلينيكية الكاملة (End-to-End Testing):")

    headers_tests = ["رقم الفحص", "السيناريو المختبر", "المدخلات والإجراء", "النتيجة المتوقعة", "الحالة الفعلية"]
    data_tests = [
        ["TC-01", "تسجيل دخول وتوجيه الصلاحيات", "دخول ببيانات دكتور / فني مختبر / صيدلي", "توليد توكن Sanctum والتوجيه للمحطة المخصصة لكل دور", "ناجح (Passed 100%)"],
        ["TC-02", "حماية الـ APIs ومنع الوصول غير المصرح", "محاولة استدعاء /api/users بواسطة حساب الصيدلي", "إرجاع كود الاستجابة 403 Forbidden ومنع الوصول", "ناجح (Passed 100%)"],
        ["TC-03", "طلب وفحص تحليل مخبري إكلينيكي", "طلب الطبيب لتحليل CBC وإدخال المخبري للقيمة 13.5", "تحديث الحالة إلى مكتمل وظهور القيمة فوراً في شاشة الطبيب", "ناجح (Passed 100%)"],
        ["TC-04", "صرف الدواء والخصم الآلي", "صرف دواء Amoxicillin بواسطة الصيدلي", "تحديث حالة الوصفة إلى مصروفة وخصم علبة من رصيد المستودع", "ناجح (Passed 100%)"],
        ["TC-05", "السجل الطبي التراكمي للمريض", "استعراض تاريخ المريض رقم #1", "عرض كافة الزيارات السابقة والتشخيصات والتحاليل والوصفات", "ناجح (Passed 100%)"]
    ]
    add_styled_table(headers_tests, data_tests)

    # 7. الفصل السادس
    add_heading_1("الفصل السادس: الخاتمة والتوصيات المستقبلية (Conclusion & Future Scope)")
    add_heading_2("6.1 النتائج المحققة (Achievements)")
    add_body_p("نجح المشروع في تحقيق كافة أهدافه المحددة، وتم بناء نظام إدارة طبي إكلينيكي متكامل ومؤتمت بالكامل يقضي على البيروقراطية الورقية ويربط كافة مفاصل المركز الطبي في بيئة برمجية آمنة، سريعة، وسهلة الاستخدام.")

    add_heading_2("6.2 الآفاق والتطويرات المستقبلية (Future Enhancements)")
    add_bullet("إضافة نماذج ذكاء اصطناعي (AI Medical Assistant) لاقتراح الفحوصات الطبية بناءً على الأعراض المدخلة والتاريخ المرضي.", "1. الذكاء الاصطناعي في المساعدة التشخيصية:")
    add_bullet("تطوير تطبيق مخصص للهواتف الذكية (Flutter Mobile App) يتيح للمريض حجز المواعيد واستعراض نتائج تحاليل ووصفاته.", "2. تطبيق المريض على الهواتف الذكية:")
    add_bullet("ربط النظام مع بوابات شركات التأمين الصحي للتحقق الفوري من التغطية المالية للمريض.", "3. الربط مع شبكات التأمين الصحي:")
    add_bullet("دعم معايير HL7 و FHIR العالمية لتبادل السجلات الطبية بين المستشفيات والمراكز الكبرى.", "4. التكامل مع المعايير الصحية العالمية:")

    # 8. المراجع
    add_heading_1("المراجع والمصادر (References)")
    add_bullet("Laravel Framework Documentation (Version 10.x) - https://laravel.com/docs")
    add_bullet("Laravel Sanctum API Authentication Guide - https://laravel.com/docs/sanctum")
    add_bullet("MySQL 8.0 Reference Manual - Relational Database Systems")
    add_bullet("World Health Organization (WHO) - Guidelines on Primary Healthcare Digital Health Systems")
    add_bullet("Role-Based Access Control (RBAC) Standards - NIST Special Publication")

    # Save
    out_path = r"c:\Users\Zaid\Desktop\clinic\clinic\Medical_Point_Management_System_Academic_Report.docx"
    doc.save(out_path)
    print(f"[SUCCESS] Academic report created at: {out_path}")

if __name__ == "__main__":
    create_document()
